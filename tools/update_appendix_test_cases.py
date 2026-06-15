from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


DOCX = Path(
    r"C:\Users\Tempest\Documents\GitHub\digital-market\artifacts\Диплом Гредягин К.С. ИС-41 2026_глава3_актуализировано.docx"
)


def para_text(element) -> str:
    return "".join(node.text or "" for node in element.xpath(".//w:t")).strip()


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=70, start=70, bottom=70, end=70) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_in: float) -> None:
    width = int(width_in * 1440)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def clear_after_appendix_2(doc: Document) -> None:
    body = doc.element.body
    children = list(body.iterchildren())
    start_idx = None
    sect_idx = None

    for idx, child in enumerate(children):
        tag = child.tag.rsplit("}", 1)[-1]
        if tag == "sectPr":
            sect_idx = idx
            break
        if tag == "p" and para_text(child) == "Приложение 2. Тест-кейсы":
            start_idx = idx

    if start_idx is None:
        raise RuntimeError("Не найдено начало 'Приложение 2. Тест-кейсы'.")

    end_idx = sect_idx if sect_idx is not None else len(children)
    for child in children[start_idx:end_idx]:
        body.remove(child)


def add_paragraph(doc: Document, text: str = "", *, bold: bool = False, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.bold = bold
    return p


def add_heading(doc: Document, text: str, level: int) -> None:
    p = doc.add_paragraph()
    p.style = f"Heading {level}" if f"Heading {level}" in {s.name for s in doc.styles} else "Normal"
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    p.add_run(text)


def add_caption(doc: Document, text: str) -> None:
    p = add_paragraph(doc, text, bold=True)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)


def add_table(doc: Document, rows: list[list[str]]) -> None:
    widths = [0.82, 1.65, 3.15, 2.25]
    table = doc.add_table(rows=0, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    try:
        table.style = "Table Grid"
    except Exception:
        pass

    for row_idx, values in enumerate(rows):
        row = table.add_row()
        for col_idx, value in enumerate(values):
            cell = row.cells[col_idx]
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_width(cell, widths[col_idx])
            set_cell_margins(cell)
            if row_idx == 0:
                set_cell_shading(cell, "E5E7EB")
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            run = paragraph.add_run(value)
            run.font.size = Pt(8)
            if row_idx == 0:
                run.bold = True
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            elif col_idx == 0:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_paragraph(doc)


TEST_SECTIONS = [
    (
        "Таблица 4. Регистрация и аутентификация - позитивные сценарии",
        [
            ["ID", "Название теста", "Шаги", "Ожидаемый результат"],
            ["AUTH-01", "Успешная регистрация", "POST /api/auth/register с email, username, password.", "Возвращён JWT-токен; пользователь создан с ролью USER и начальным балансом 1000."],
            ["AUTH-02", "Успешный вход", "POST /api/auth/login с зарегистрированными email и password.", "Возвращён JWT-токен, пригодный для защищённых запросов."],
            ["AUTH-03", "Получение текущего пользователя", "GET /api/auth/me с Authorization: Bearer <token>.", "Возвращены id, email, username, role и balance текущего пользователя."],
            ["AUTH-04", "Смена username", "PATCH /api/auth/me с новым username и токеном.", "Username обновлён, в ответе возвращены актуальные данные пользователя."],
        ],
    ),
    (
        "Таблица 5. Регистрация и аутентификация - негативные сценарии",
        [
            ["ID", "Название теста", "Шаги", "Ожидаемый результат"],
            ["AUTH-N01", "Повтор email", "POST /api/auth/register с email уже существующего пользователя.", "Возвращается ошибка конфликта, новый пользователь не создаётся."],
            ["AUTH-N02", "Повтор username", "POST /api/auth/register с занятым username.", "Возвращается ошибка конфликта, новый пользователь не создаётся."],
            ["AUTH-N03", "Неверный пароль", "POST /api/auth/login с валидным email и неверным password.", "Возвращается ошибка авторизации, токен не выдаётся."],
            ["AUTH-N04", "Запрос /me без токена", "GET /api/auth/me без Authorization.", "Возвращается 401, данные пользователя не раскрываются."],
            ["AUTH-N05", "Некорректный JWT", "GET /api/auth/me с повреждённым или просроченным токеном.", "Возвращается 401 Invalid token."],
            ["AUTH-N06", "Смена username на занятый", "PATCH /api/auth/me с username другого пользователя.", "Возвращается ошибка конфликта, данные не изменяются."],
        ],
    ),
    (
        "Таблица 6. Управление товарами и каталог - позитивные сценарии",
        [
            ["ID", "Название теста", "Шаги", "Ожидаемый результат"],
            ["ITEM-01", "Получение каталога", "GET /api/items без параметров.", "Возвращаются товары со статусом LISTED, цена преобразована в число, есть hasImage и рейтинг продавца."],
            ["ITEM-02", "Поиск по названию", "GET /api/items?search=<строка>.", "Возвращаются товары, название которых соответствует поиску без учёта регистра."],
            ["ITEM-03", "Фильтр по категории", "GET /api/items?category=<CATEGORY>.", "Возвращаются товары выбранной категории."],
            ["ITEM-04", "Фильтр по статусу", "GET /api/items?status=ALL.", "Возвращаются товары разных статусов, если они есть в БД."],
            ["ITEM-05", "Карточка товара", "GET /api/items/:id.", "Возвращены детальные данные товара, продавец, владелец при наличии и рейтинг продавца."],
            ["ITEM-06", "Получение изображения", "GET /api/items/:id/image для товара с imageData.", "Возвращаются бинарные данные изображения с корректным Content-Type."],
            ["ITEM-07", "Мои товары", "GET /api/items/mine с токеном продавца.", "Возвращается список товаров текущего пользователя."],
            ["ITEM-08", "Создание товара без изображения", "POST /api/items multipart/form-data без image.", "Товар создан со статусом LISTED и категорией, hasImage=false."],
            ["ITEM-09", "Создание товара с изображением", "POST /api/items multipart/form-data с jpg/png/webp до 5 MB.", "Товар создан, imageData/imageMime сохранены, hasImage=true."],
            ["ITEM-10", "Редактирование JSON", "PUT /api/items/:id с title/description/price/category.", "Данные непроданного товара обновлены владельцем."],
            ["ITEM-11", "Замена изображения", "PATCH /api/items/:id multipart/form-data с новым image.", "Изображение заменено, hasImage=true."],
            ["ITEM-12", "Удаление изображения", "PATCH /api/items/:id с removeImage=true.", "imageData и imageMime очищены, hasImage=false."],
            ["ITEM-13", "Удаление товара", "DELETE /api/items/:id для своего непроданного товара.", "Товар удалён, повторное получение возвращает ошибку."],
        ],
    ),
    (
        "Таблица 7. Управление товарами и каталог - негативные сценарии",
        [
            ["ID", "Название теста", "Шаги", "Ожидаемый результат"],
            ["ITEM-N01", "Создание без авторизации", "POST /api/items без токена.", "Возвращается 401, товар не создаётся."],
            ["ITEM-N02", "Редактирование чужого товара", "PUT /api/items/:id с токеном другого пользователя.", "Возвращается 403, данные не изменяются."],
            ["ITEM-N03", "Удаление чужого товара", "DELETE /api/items/:id с токеном другого пользователя.", "Возвращается 403, товар не удаляется."],
            ["ITEM-N04", "Редактирование проданного товара", "PUT или PATCH /api/items/:id для товара SOLD.", "Возвращается ошибка Cannot edit sold item."],
            ["ITEM-N05", "Удаление проданного товара", "DELETE /api/items/:id для товара SOLD.", "Возвращается ошибка Cannot delete sold item."],
            ["ITEM-N06", "Некорректная категория", "POST/PATCH /api/items с category не из enum.", "Возвращается ошибка валидации, данные не сохраняются."],
            ["ITEM-N07", "Некорректная цена", "POST/PATCH /api/items с price <= 0 или не числом.", "Возвращается ошибка валидации."],
            ["ITEM-N08", "Неподдерживаемый MIME", "POST/PATCH /api/items с image не jpg/png/webp.", "Возвращается ошибка, изображение не принимается."],
            ["ITEM-N09", "Файл больше лимита", "POST/PATCH /api/items с image больше 5 MB.", "Запрос отклоняется, изображение не сохраняется."],
            ["ITEM-N10", "Изображение отсутствует", "GET /api/items/:id/image для товара без изображения.", "Возвращается 404."],
            ["ITEM-N11", "Несуществующий товар", "GET /api/items/:id для несуществующего id.", "Возвращается 404 Item not found."],
        ],
    ),
    (
        "Таблица 8. Заказы, платежи и сделки - позитивные сценарии",
        [
            ["ID", "Название теста", "Шаги", "Ожидаемый результат"],
            ["ORD-01", "Покупка с баланса", "POST /api/orders/purchase/:itemId с paymentMethod=BALANCE.", "Созданы Order и Payment; товар SOLD, ownerId=покупатель; баланс покупателя уменьшен; создана WalletTransaction PURCHASE_HOLD."],
            ["ORD-02", "Покупка картой", "POST /api/orders/purchase/:itemId с paymentMethod=CARD.", "Создан Payment method=CARD, provider=mock-card-acquiring; внутренний баланс покупателя при покупке не уменьшается."],
            ["ORD-03", "Покупка через СБП", "POST /api/orders/purchase/:itemId с paymentMethod=SBP.", "Создан Payment method=SBP, provider=mock-sbp; сделка получает статус PAID."],
            ["ORD-04", "Детальная сделка", "GET /api/orders/:orderId участником сделки.", "Возвращены item, buyer, seller, payment, messages, review и нормализованные суммы."],
            ["ORD-05", "История покупок", "GET /api/orders/my с токеном покупателя.", "Возвращены сделки, где пользователь является buyer."],
            ["ORD-06", "История продаж", "GET /api/orders/sales с токеном продавца.", "Возвращены сделки, где пользователь является seller."],
            ["ORD-07", "Сообщение в чате сделки", "POST /api/orders/:orderId/messages с body.", "Создано сообщение типа TEXT с автором."],
            ["ORD-08", "Выдача оплаченных данных", "PATCH /api/orders/:orderId/delivery продавцом.", "deliveryData сохранены; добавлены сообщения DELIVERY и SYSTEM."],
            ["ORD-09", "Подтверждение получения", "POST /api/orders/:orderId/confirm покупателем.", "Order=COMPLETED; Payment=RELEASED; confirmedAt заполнен; продавцу начислены средства; создана WalletTransaction SALE_RELEASE."],
            ["ORD-10", "Возврат до выдачи данных", "POST /api/orders/:orderId/refund покупателем до deliveryData.", "Order=CANCELLED; Payment=REFUNDED; товар снова LISTED; ownerId очищен; при BALANCE деньги возвращены покупателю."],
            ["ORD-11", "Открытие спора", "POST /api/orders/:orderId/dispute участником сделки.", "Order=DISPUTED, disputedAt заполнен, в чат добавлено системное сообщение."],
            ["ORD-12", "Возврат после спора", "После deliveryData открыть спор, затем POST /refund.", "Сделка отменена, платёж помечен REFUNDED, товар возвращён в каталог по правилам возврата."],
            ["ORD-13", "Повторное подтверждение завершённой сделки", "Повторить POST /confirm для уже COMPLETED сделки.", "Возвращается текущее состояние сделки без повторного начисления средств."],
        ],
    ),
    (
        "Таблица 9. Заказы, платежи и сделки - негативные сценарии",
        [
            ["ID", "Название теста", "Шаги", "Ожидаемый результат"],
            ["ORD-N01", "Покупка без авторизации", "POST /api/orders/purchase/:itemId без токена.", "Возвращается 401, заказ не создаётся."],
            ["ORD-N02", "Покупка собственного товара", "Продавец вызывает POST /api/orders/purchase/:itemId для своего товара.", "Возвращается ошибка Seller cannot buy own item."],
            ["ORD-N03", "Покупка проданного товара", "Повторить покупку товара со статусом SOLD.", "Возвращается ошибка Item is not available, второй заказ не создаётся."],
            ["ORD-N04", "Недостаточно средств", "paymentMethod=BALANCE при balance меньше цены.", "Возвращается ошибка Not enough balance, товар остаётся LISTED."],
            ["ORD-N05", "Чужая сделка", "GET /api/orders/:orderId пользователем, который не buyer и не seller.", "Возвращается 403 Forbidden."],
            ["ORD-N06", "Сообщение в чужую сделку", "POST /messages пользователем вне сделки.", "Возвращается 403, сообщение не создаётся."],
            ["ORD-N07", "Выдача данных покупателем", "PATCH /delivery от buyer вместо seller.", "Возвращается 403 Only seller can add delivery data."],
            ["ORD-N08", "Подтверждение продавцом", "POST /confirm от seller.", "Возвращается 403 Only buyer can confirm delivery."],
            ["ORD-N09", "Возврат продавцом", "POST /refund от seller.", "Возвращается 403 Only buyer can request refund."],
            ["ORD-N10", "Возврат завершённой сделки", "POST /refund после статуса COMPLETED.", "Возвращается ошибка Completed deal cannot be refunded."],
            ["ORD-N11", "Возврат после выдачи без спора", "PATCH /delivery, затем POST /refund без POST /dispute.", "Возвращается ошибка Open a dispute before refunding delivered data."],
            ["ORD-N12", "Спор по завершённой или отменённой сделке", "POST /dispute для COMPLETED или CANCELLED.", "Возвращается ошибка, статус не меняется."],
            ["ORD-N13", "Несуществующая сделка", "GET /api/orders/:orderId с несуществующим id.", "Возвращается 404 Order not found."],
        ],
    ),
    (
        "Таблица 10. Отзывы, рейтинг и профиль продавца",
        [
            ["ID", "Название теста", "Шаги", "Ожидаемый результат"],
            ["REV-01", "Создание отзыва", "После COMPLETED выполнить POST /api/orders/:orderId/review с rating и comment.", "Отзыв создан, связан с orderId, buyerId и sellerId."],
            ["REV-02", "Обновление отзыва", "Повторить POST /review по той же сделке с новой оценкой.", "Сработал upsert: отзыв обновлён, дубль не создан."],
            ["REV-03", "Расчёт рейтинга продавца", "GET /api/users/:userId после нескольких отзывов.", "Возвращены average и count по отзывам продавца."],
            ["REV-04", "Рейтинг в каталоге", "GET /api/items и GET /api/items/:id.", "В данных seller присутствует rating, который можно показать в интерфейсе."],
            ["USR-01", "Публичный профиль продавца", "GET /api/users/:userId.", "Возвращены user, rating, активные товары и последние отзывы."],
            ["REV-N01", "Отзыв до завершения", "POST /review для сделки PAID или DISPUTED.", "Возвращается ошибка Only completed deals can be reviewed."],
            ["REV-N02", "Отзыв от продавца", "POST /review от seller.", "Возвращается 403 Only buyer can review this deal."],
            ["REV-N03", "Оценка вне диапазона", "POST /review с rating меньше 1 или больше 5.", "Возвращается ошибка валидации."],
            ["USR-N01", "Несуществующий продавец", "GET /api/users/:userId для несуществующего пользователя.", "Возвращается 404 User not found."],
        ],
    ),
    (
        "Таблица 11. Личная переписка с продавцом",
        [
            ["ID", "Название теста", "Шаги", "Ожидаемый результат"],
            ["MSG-01", "Отправка личного сообщения", "POST /api/users/:userId/messages с token и body.", "Создано DirectMessage с sender, recipient и body."],
            ["MSG-02", "Получение переписки", "GET /api/users/:userId/messages с token.", "Возвращается переписка двух пользователей в порядке createdAt asc."],
            ["MSG-N01", "Переписка без токена", "GET или POST /messages без Authorization.", "Возвращается 401."],
            ["MSG-N02", "Сообщение самому себе", "POST /api/users/<свой id>/messages.", "Возвращается ошибка Cannot message yourself."],
            ["MSG-N03", "Пустое сообщение", "POST /messages с пустым body.", "Возвращается ошибка валидации."],
            ["MSG-N04", "Несуществующий адресат", "POST /messages для несуществующего userId.", "Возвращается 404 User not found."],
        ],
    ),
    (
        "Таблица 12. Клиентский интерфейс - пользовательские сценарии",
        [
            ["ID", "Название теста", "Шаги", "Ожидаемый результат"],
            ["UI-01", "Каталог", "Открыть /, выполнить поиск и выбрать категорию.", "Список товаров обновляется, пустой результат отображается корректно."],
            ["UI-02", "Карточка товара", "Открыть /items/:id, выбрать способ оплаты.", "Показываются данные товара, рейтинг продавца и доступные варианты оплаты."],
            ["UI-03", "Недостаточный баланс в UI", "Открыть товар дороже баланса и выбрать оплату BALANCE.", "Интерфейс предупреждает пользователя или переключает доступный вариант оплаты."],
            ["UI-04", "Переход после покупки", "Купить товар из карточки.", "После успешного запроса пользователь попадает на страницу созданной сделки, баланс обновляется через refreshMe."],
            ["UI-05", "Страница сделки", "Открыть /deals/:id для покупателя и продавца.", "Отображаются товар, платёж, статус, чат и действия, доступные роли пользователя."],
            ["UI-06", "Профиль", "Открыть /profile и проверить покупки, продажи, товары и настройки.", "Вкладки отображают актуальные данные пользователя."],
            ["UI-07", "Профиль продавца", "Открыть /users/:id.", "Отображаются рейтинг, отзывы, активные товары и форма личного сообщения."],
            ["UI-08", "Локализация и валюта", "В настройках переключить ru/en и RUB/USD.", "Тексты интерфейса и формат сумм меняются согласно SettingsContext."],
        ],
    ),
]


def main() -> None:
    doc = Document(DOCX)
    clear_after_appendix_2(doc)

    add_heading(doc, "Приложение 2. Тест-кейсы", 1)
    add_paragraph(
        doc,
        "Ниже приведён набор тест-кейсов для проверки фактически реализованных модулей маркетплейса: авторизации, каталога, сделок, платежей, операций баланса, отзывов, профиля продавца, личной переписки и пользовательского интерфейса.",
    )

    for caption, rows in TEST_SECTIONS:
        add_caption(doc, caption)
        add_table(doc, rows)

    doc.save(DOCX)
    print(DOCX)


if __name__ == "__main__":
    main()
