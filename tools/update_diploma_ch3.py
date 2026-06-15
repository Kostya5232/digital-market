from __future__ import annotations

from pathlib import Path
from textwrap import dedent

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = Path(r"C:\Users\Tempest\Desktop\Диплом\Диплом Гредягин К.С. ИС-41 2026_дополнено.docx")
OUT_DIR = ROOT / "artifacts"
OUT_DOCX = OUT_DIR / "Диплом Гредягин К.С. ИС-41 2026_глава3_актуализировано.docx"
ER_IMAGE = OUT_DIR / "er_model_actual.png"


def load_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    candidates = [
        Path(r"C:\Windows\Fonts\arialbd.ttf" if bold else r"C:\Windows\Fonts\arial.ttf"),
        Path(r"C:\Windows\Fonts\calibrib.ttf" if bold else r"C:\Windows\Fonts\calibri.ttf"),
    ]
    for path in candidates:
        if path.exists():
            return ImageFont.truetype(str(path), size)
    return ImageFont.load_default()


def wrap_text(draw: ImageDraw.ImageDraw, text: str, font: ImageFont.FreeTypeFont, width: int) -> list[str]:
    lines: list[str] = []
    for raw in text.splitlines():
        words = raw.split()
        line = ""
        for word in words:
            candidate = word if not line else f"{line} {word}"
            if draw.textlength(candidate, font=font) <= width:
                line = candidate
            else:
                if line:
                    lines.append(line)
                line = word
        lines.append(line)
    return lines


def draw_entity(
    draw: ImageDraw.ImageDraw,
    box: tuple[int, int, int, int],
    title: str,
    fields: list[str],
    fill: str,
    accent: str,
) -> None:
    x1, y1, x2, y2 = box
    draw.rounded_rectangle(box, radius=18, fill=fill, outline=accent, width=3)
    draw.rounded_rectangle((x1, y1, x2, y1 + 44), radius=18, fill=accent, outline=accent)
    draw.rectangle((x1, y1 + 24, x2, y1 + 44), fill=accent)

    title_font = load_font(24, bold=True)
    body_font = load_font(18)
    draw.text((x1 + 18, y1 + 9), title, fill="white", font=title_font)

    y = y1 + 60
    for field in fields:
        for line in wrap_text(draw, field, body_font, x2 - x1 - 36):
            draw.text((x1 + 18, y), line, fill="#1f2937", font=body_font)
            y += 25
        y += 2


def draw_link(draw: ImageDraw.ImageDraw, start: tuple[int, int], end: tuple[int, int], label: str = "") -> None:
    color = "#6b7280"
    draw.line((start, end), fill=color, width=3)
    if label:
        font = load_font(16)
        mx = (start[0] + end[0]) // 2
        my = (start[1] + end[1]) // 2
        bbox = draw.textbbox((mx, my), label, font=font)
        pad = 4
        draw.rounded_rectangle(
            (bbox[0] - pad, bbox[1] - pad, bbox[2] + pad, bbox[3] + pad),
            radius=6,
            fill="#f8fafc",
            outline="#d1d5db",
        )
        draw.text((mx, my), label, fill="#374151", font=font)


def draw_poly_link(draw: ImageDraw.ImageDraw, points: list[tuple[int, int]], label_pos: tuple[int, int], label: str) -> None:
    color = "#6b7280"
    for a, b in zip(points, points[1:]):
        draw.line((a, b), fill=color, width=3)
    font = load_font(16)
    x, y = label_pos
    bbox = draw.textbbox((x, y), label, font=font)
    pad = 4
    draw.rounded_rectangle(
        (bbox[0] - pad, bbox[1] - pad, bbox[2] + pad, bbox[3] + pad),
        radius=6,
        fill="#f8fafc",
        outline="#d1d5db",
    )
    draw.text((x, y), label, fill="#374151", font=font)


def make_er_diagram(path: Path) -> None:
    img = Image.new("RGB", (1800, 1160), "#f8fafc")
    draw = ImageDraw.Draw(img)
    title_font = load_font(34, bold=True)
    subtitle_font = load_font(20)
    draw.text((60, 35), "ER-модель реализованной базы данных digital_market", fill="#111827", font=title_font)
    draw.text(
        (60, 80),
        "Схема отражает фактические модели Prisma: пользователи, товары, сделки, платежи, кошелёк, сообщения и отзывы.",
        fill="#4b5563",
        font=subtitle_font,
    )

    boxes = {
        "User": (70, 160, 420, 380),
        "Item": (720, 160, 1070, 390),
        "DirectMessage": (1380, 160, 1730, 360),
        "Order": (720, 505, 1070, 780),
        "OrderMessage": (70, 620, 420, 830),
        "Review": (1380, 600, 1730, 810),
        "Payment": (595, 910, 945, 1095),
        "WalletTransaction": (1110, 900, 1530, 1095),
    }

    palette = {
        "User": ("#e0f2fe", "#0369a1"),
        "Item": ("#ecfdf5", "#047857"),
        "DirectMessage": ("#fef3c7", "#b45309"),
        "Order": ("#eef2ff", "#4338ca"),
        "OrderMessage": ("#f3f4f6", "#4b5563"),
        "Review": ("#fce7f3", "#be185d"),
        "Payment": ("#ede9fe", "#6d28d9"),
        "WalletTransaction": ("#dcfce7", "#15803d"),
    }

    entity_fields = {
        "User": [
            "id, email, username, passwordHash",
            "role, balance",
            "items, orders, sales",
            "reviews, messages",
        ],
        "Item": [
            "id, title, description, price",
            "status, category",
            "imageData, imageMime",
            "sellerId, ownerId",
        ],
        "DirectMessage": [
            "senderId -> User.id",
            "recipientId -> User.id",
            "body, createdAt",
        ],
        "Order": [
            "itemId, buyerId, sellerId",
            "price, status",
            "deliveryData",
            "confirmedAt, disputedAt, cancelledAt",
        ],
        "OrderMessage": [
            "orderId -> Order.id",
            "authorId -> User.id",
            "body, type, createdAt",
        ],
        "Review": [
            "orderId UNIQUE",
            "buyerId, sellerId",
            "rating, comment",
        ],
        "Payment": [
            "orderId UNIQUE",
            "method, status, amount",
            "provider, releasedAt, refundedAt",
        ],
        "WalletTransaction": [
            "userId, orderId, paymentId",
            "type, amount",
            "balanceAfter, note",
        ],
    }

    for name, box in boxes.items():
        draw_entity(draw, box, name, entity_fields[name], *palette[name])

    draw_link(draw, (420, 260), (720, 260), "seller/owner")
    draw_link(draw, (895, 390), (895, 505), "item")
    draw_link(draw, (720, 635), (420, 720), "messages")
    draw_link(draw, (1070, 650), (1380, 690), "review")
    draw_link(draw, (855, 780), (760, 910), "payment")
    draw_link(draw, (1070, 735), (1110, 980), "wallet")
    draw_link(draw, (420, 335), (720, 570), "buyer/seller")
    draw_poly_link(draw, [(245, 160), (245, 125), (1555, 125), (1555, 160)], (1435, 116), "direct chat")

    path.parent.mkdir(parents=True, exist_ok=True)
    img.save(path)


def find_style(doc: Document, preferred: str, fallback: str = "Normal") -> str:
    names = {style.name for style in doc.styles}
    return preferred if preferred in names else fallback


def ensure_code_style(doc: Document) -> str:
    name = "Code Block"
    if name not in {style.name for style in doc.styles}:
        style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        style.base_style = doc.styles["Normal"]
        style.font.name = "Courier New"
        style.font.size = Pt(8.5)
        style.font.color.rgb = RGBColor(31, 41, 55)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.line_spacing = 1.0
        style.paragraph_format.left_indent = Inches(0.18)
    return name


def shade_paragraph(paragraph, fill: str = "F3F4F6") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_in: float) -> None:
    width = int(width_in * 1440)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=80, start=80, bottom=80, end=80) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def move_before_anchor(doc: Document, element, anchor) -> None:
    body = doc.element.body
    body.remove(element)
    anchor.addprevious(element)


class ChapterWriter:
    def __init__(self, doc: Document, anchor):
        self.doc = doc
        self.anchor = anchor
        self.body_style = find_style(doc, "Обычный текст", "Normal")
        self.code_style = ensure_code_style(doc)

    def paragraph(self, text: str = "", style: str | None = None, *, bold: bool = False, italic: bool = False):
        p = self.doc.add_paragraph()
        move_before_anchor(self.doc, p._element, self.anchor)
        if style:
            p.style = find_style(self.doc, style, self.body_style)
        else:
            p.style = self.body_style
        if text:
            run = p.add_run(text)
            run.bold = bold
            run.italic = italic
        return p

    def heading(self, text: str, level: int):
        return self.paragraph(text, f"Heading {level}")

    def caption(self, text: str):
        p = self.paragraph(text, self.body_style, bold=True)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(3)
        return p

    def code(self, code: str):
        for line in dedent(code).strip("\n").splitlines():
            p = self.paragraph(line.rstrip(), self.code_style)
            shade_paragraph(p)

    def picture(self, path: Path, width: float = 6.35):
        p = self.paragraph("", "Normal")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(path), width=Inches(width))
        return p

    def table(self, rows: list[list[str]], widths: list[float]):
        table = self.doc.add_table(rows=0, cols=len(widths))
        move_before_anchor(self.doc, table._element, self.anchor)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        try:
            table.style = "Table Grid"
        except Exception:
            pass

        for row_idx, row_data in enumerate(rows):
            row = table.add_row()
            for col_idx, text in enumerate(row_data):
                cell = row.cells[col_idx]
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                set_cell_width(cell, widths[col_idx])
                set_cell_margins(cell)
                if row_idx == 0:
                    set_cell_shading(cell, "E5E7EB")
                para = cell.paragraphs[0]
                para.style = self.body_style
                para.paragraph_format.space_after = Pt(0)
                if text:
                    run = para.add_run(text)
                    run.font.size = Pt(9)
                    if row_idx == 0:
                        run.bold = True
                if col_idx > 0 and row_idx > 0 and len(text) < 22:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.paragraph("")
        return table


def remove_old_chapter(doc: Document):
    body = doc.element.body
    children = list(body.iterchildren())

    def para_text(el) -> str:
        return "".join(node.text or "" for node in el.xpath(".//w:t")).strip()

    start_idx = None
    end_idx = None
    for idx, child in enumerate(children):
        tag = child.tag.rsplit("}", 1)[-1]
        if tag != "p":
            continue
        text = para_text(child)
        if text == "Глава 3. Реализация":
            start_idx = idx
        elif start_idx is not None and text.lower().startswith("заключение"):
            end_idx = idx
            break

    if start_idx is None or end_idx is None:
        raise RuntimeError("Не удалось найти границы главы 3.")

    anchor = children[end_idx]
    for child in children[start_idx:end_idx]:
        body.remove(child)
    return anchor


def build_chapter(writer: ChapterWriter) -> None:
    writer.heading("Глава 3. Реализация", 1)

    writer.heading("3.1. Выбор технологий и инструментов", 2)
    writer.paragraph(
        "Реализация веб-приложения выполнена как полнофункциональный клиент-серверный прототип маркетплейса цифровых товаров. "
        "Технологический стек выбран с учётом требований к интерактивному интерфейсу, безопасной авторизации, транзакционному оформлению сделок и удобному сопровождению кода."
    )
    writer.table(
        [
            ["Уровень", "Использованные технологии", "Назначение в проекте"],
            ["Клиент", "React 19, TypeScript, Vite, React Router", "Построение SPA-интерфейса, маршрутизация, компонентная структура страниц и форм."],
            ["Сервер", "Node.js, Express, TypeScript", "REST API, бизнес-логика покупок, сделок, отзывов, чатов и профилей."],
            ["Данные", "PostgreSQL, Prisma ORM", "Реляционная модель предметной области, миграции, транзакции, типизированный доступ к данным."],
            ["Безопасность", "JWT, bcryptjs, middleware requireAuth", "Хэширование паролей, выдача токена, проверка защищённых действий пользователя."],
            ["Валидация и файлы", "zod, multer", "Проверка входных данных и загрузка изображений товаров в память перед сохранением в БД."],
        ],
        [1.25, 2.2, 3.05],
    )
    writer.paragraph(
        "Клиентская часть размещена в каталоге client и запускается через Vite. Серверная часть размещена в каталоге server, использует Express-приложение и Prisma Client для доступа к PostgreSQL. "
        "Такое разделение позволяет независимо развивать интерфейс и API, сохраняя единую структуру репозитория."
    )

    writer.heading("3.2. Архитектура программной системы", 2)
    writer.paragraph(
        "Система построена по классической схеме: пользователь работает с React-интерфейсом, интерфейс отправляет HTTP-запросы к REST API с префиксом /api, сервер проверяет данные и права доступа, затем выполняет операции в PostgreSQL через Prisma. "
        "Ответы API возвращаются в формате JSON, а изображения товаров отдаются отдельным бинарным эндпоинтом."
    )

    writer.heading("3.2.1. Общая схема взаимодействия компонентов", 3)
    writer.paragraph(
        "Основной поток данных выглядит следующим образом: React UI -> API-модуль клиента -> Express route -> middleware авторизации/валидации -> Prisma Client -> PostgreSQL. "
        "Например, при покупке товара страница карточки вызывает клиентскую функцию buyItem, сервер создаёт сделку, платёж и системное сообщение в одной транзакции, после чего интерфейс перенаправляет пользователя на страницу сделки."
    )

    writer.heading("3.2.2. Серверное приложение и основные маршруты", 3)
    writer.table(
        [
            ["Модуль", "Маршруты", "Реализованная логика"],
            ["Auth", "POST /api/auth/register\nPOST /api/auth/login\nGET /api/auth/me\nPATCH /api/auth/me", "Регистрация, вход, получение текущего пользователя, изменение username с проверкой уникальности."],
            ["Items", "GET /api/items\nGET /api/items/mine\nGET /api/items/:id\nGET /api/items/:id/image\nPOST /api/items\nPUT/PATCH /api/items/:id\nDELETE /api/items/:id", "Каталог, поиск, фильтрация, карточка товара, товары продавца, загрузка/замена изображения, редактирование и удаление."],
            ["Orders", "POST /api/orders/purchase/:itemId\nGET /api/orders/my\nGET /api/orders/sales\nGET /api/orders/:orderId\nPOST /messages\nPATCH /delivery\nPOST /confirm\nPOST /refund\nPOST /dispute\nPOST /review", "Создание сделки, история покупок и продаж, чат сделки, выдача данных, подтверждение, возврат, спор и отзыв."],
            ["Users", "GET /api/users/:userId\nGET /api/users/:userId/messages\nPOST /api/users/:userId/messages", "Публичный профиль продавца, рейтинг, отзывы, активные товары и личная переписка."],
        ],
        [1.15, 2.6, 2.75],
    )
    writer.paragraph(
        "Доступ к защищённым маршрутам ограничен middleware requireAuth. Токен передаётся в заголовке Authorization: Bearer <token>, а ошибки обрабатываются единым middleware errorHandler, что делает ответы API единообразными."
    )

    writer.heading("3.2.3. Клиентское приложение", 3)
    writer.paragraph(
        "Клиент реализован как одностраничное приложение. Маршрутизация выполняется через BrowserRouter, а доступные экраны определены в App.tsx: каталог, регистрация, вход, профиль, список сделок, карточка сделки, публичный профиль продавца, добавление и редактирование товара."
    )
    writer.table(
        [
            ["Экран", "Маршрут", "Назначение"],
            ["Home", "/", "Каталог товаров, поиск по названию, фильтрация по категориям, отображение баланса и статистики каталога."],
            ["ItemDetail", "/items/:id", "Карточка товара, выбор способа оплаты, создание сделки, переход к профилю продавца."],
            ["Profile", "/profile", "Товары пользователя, покупки, продажи, настройки языка, валюты и username."],
            ["Deals", "/deals", "Раздел покупок и продаж с фильтрацией по роли пользователя в сделке."],
            ["DealDetail", "/deals/:id", "Детали сделки, оплата, выдача данных, чат, подтверждение, возврат, спор и отзыв."],
            ["SellerProfile", "/users/:id", "Публичная страница продавца, рейтинг, отзывы, активные товары и личный чат."],
        ],
        [1.35, 1.45, 3.7],
    )
    writer.paragraph(
        "Общее состояние авторизации вынесено в AuthContext: токен хранится в localStorage, данные пользователя загружаются через GET /api/auth/me, а метод refreshMe используется после операций, меняющих баланс. "
        "Настройки языка и валюты вынесены в SettingsContext; интерфейс поддерживает ru/en и отображение сумм в RUB/USD."
    )

    writer.heading("3.3. Проектирование базы данных", 2)
    writer.paragraph(
        "База данных спроектирована вокруг жизненного цикла цифрового товара: пользователь публикует товар, покупатель оформляет сделку, платёж резервируется, продавец передаёт оплаченные данные, покупатель подтверждает получение или открывает спор. "
        "Фактическая структура описана в schema.prisma и включает не только пользователей, товары и заказы, но также платежи, транзакции кошелька, сообщения сделки, отзывы и личные сообщения."
    )

    writer.heading("3.3.1. ER-модель и связи между сущностями", 3)
    writer.picture(ER_IMAGE)
    caption = writer.paragraph("Рисунок 1. ER-модель реализованной базы данных маркетплейса цифровых товаров.", "Normal")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    writer.paragraph(
        "Ключевые связи: пользователь может быть продавцом и покупателем в разных сделках; товар связан с продавцом и после покупки получает владельца; заказ связывает товар, покупателя и продавца; платёж относится к конкретной сделке; операции кошелька фиксируют резерв, возврат и зачисление средств; сообщения и отзывы привязаны к сделке."
    )

    writer.heading("3.3.2. Описание сущностей", 3)
    writer.table(
        [
            ["Сущность", "Ключевые поля", "Назначение"],
            ["User", "id, email, username, passwordHash, role, balance", "Пользователь системы. Может размещать товары, покупать, продавать, писать сообщения и получать отзывы."],
            ["Item", "title, description, price, status, category, imageData, imageMime, sellerId, ownerId", "Цифровой товар с категорией, изображением, продавцом и владельцем после покупки."],
            ["Order", "itemId, buyerId, sellerId, price, status, deliveryData, confirmedAt, disputedAt, cancelledAt", "Сделка между покупателем и продавцом. Хранит состояние выдачи данных и жизненный цикл покупки."],
            ["Payment", "orderId, buyerId, sellerId, method, status, amount, provider, providerPaymentId", "Платёж по сделке. Поддерживает баланс сайта, банковскую карту и СБП в демонстрационном режиме."],
            ["WalletTransaction", "userId, orderId, paymentId, type, amount, balanceAfter, note", "История движения внутреннего баланса: резерв покупки, возврат и зачисление продавцу."],
            ["OrderMessage", "orderId, authorId, body, type, createdAt", "Сообщения внутри сделки: обычные сообщения, системные события и выдача оплаченных данных."],
            ["Review", "orderId, buyerId, sellerId, rating, comment", "Отзыв покупателя о продавце по завершённой сделке. orderId уникален."],
            ["DirectMessage", "senderId, recipientId, body, createdAt", "Личная переписка пользователя с продавцом вне конкретной сделки."],
        ],
        [1.35, 2.55, 2.6],
    )

    writer.heading("3.3.3. Перечисления и статусы", 3)
    writer.table(
        [
            ["Enum", "Значения", "Использование"],
            ["Role", "USER, ADMIN", "Роль пользователя."],
            ["ItemStatus", "LISTED, SOLD", "Доступность товара в каталоге."],
            ["ItemCategory", "ACCOUNTS, KEYS, SUBSCRIPTIONS, SERVICES, GAME_CURRENCIES, NFT_TOKENS, OTHER", "Категории цифровых товаров."],
            ["OrderStatus", "PAID, COMPLETED, DISPUTED, CANCELLED", "Состояние сделки от оплаты до завершения, спора или отмены."],
            ["PaymentMethod", "BALANCE, CARD, SBP", "Выбранный способ оплаты."],
            ["PaymentStatus", "PENDING, PAID, RELEASED, REFUNDED, FAILED", "Состояние платежа."],
            ["WalletTransactionType", "PURCHASE_HOLD, REFUND, SALE_RELEASE", "Тип операции внутреннего баланса."],
            ["OrderMessageType", "TEXT, SYSTEM, DELIVERY", "Тип сообщения в чате сделки."],
        ],
        [1.55, 3.0, 1.95],
    )

    writer.caption("Листинг 1. Фрагмент схемы Prisma для платежей и кошелька")
    writer.code(
        """
        enum PaymentMethod {
          BALANCE
          CARD
          SBP
        }

        enum PaymentStatus {
          PENDING
          PAID
          RELEASED
          REFUNDED
          FAILED
        }

        enum WalletTransactionType {
          PURCHASE_HOLD
          REFUND
          SALE_RELEASE
        }

        model Payment {
          id        String        @id @default(cuid())
          orderId   String        @unique
          buyerId   String
          sellerId  String
          method    PaymentMethod
          status    PaymentStatus @default(PENDING)
          amount    Decimal       @db.Decimal(10, 2)
          provider  String?
        }
        """
    )

    writer.heading("3.4. Реализация функциональных модулей", 2)
    writer.paragraph(
        "Функциональность проекта разделена на серверные REST-модули и клиентские страницы. Ниже описаны реализованные модули с привязкой к фактическим файлам проекта."
    )

    writer.heading("3.4.1. Модуль регистрации и аутентификации", 3)
    writer.paragraph(
        "Модуль auth.ts реализует регистрацию, вход, получение текущего пользователя и изменение username. При регистрации email, username и пароль проверяются через zod, пароль хэшируется bcrypt, после чего создаётся пользователь с начальным балансом. "
        "При входе пароль сравнивается с хэшем, а клиент получает JWT-токен сроком действия 7 дней."
    )
    writer.caption("Листинг 2. Проверка JWT-токена в middleware requireAuth")
    writer.code(
        """
        export const requireAuth = (req, res, next) => {
          const header = req.headers.authorization;
          if (!header?.startsWith("Bearer ")) {
            return res.status(401).json({ message: "Missing token" });
          }

          try {
            const payload = jwt.verify(header.slice(7), env.JWT_SECRET);
            req.user = payload;
            next();
          } catch {
            res.status(401).json({ message: "Invalid token" });
          }
        };
        """
    )

    writer.heading("3.4.2. Модуль каталога и карточек товаров", 3)
    writer.paragraph(
        "Модуль items.ts отвечает за каталог, карточку товара, публикацию, редактирование и удаление. Список товаров поддерживает параметры status, search и category; по умолчанию отдаются только товары со статусом LISTED. "
        "При создании и multipart-редактировании используется multer.memoryStorage, допускаются изображения jpg/png/webp размером до 5 MB. Изображение хранится в базе как imageData и imageMime, а в API возвращается флаг hasImage."
    )
    writer.paragraph(
        "Для продавца реализованы GET /api/items/mine, POST /api/items, PUT /api/items/:id, PATCH /api/items/:id и DELETE /api/items/:id. Перед изменением или удалением сервер проверяет, что товар принадлежит текущему пользователю и ещё не продан."
    )

    writer.heading("3.4.3. Модуль сделок, платежей и арбитража", 3)
    writer.paragraph(
        "Модуль orders.ts является центральным для бизнес-логики маркетплейса. Он оформляет покупку, создаёт сделку и платёж, резервирует средства при оплате с баланса, фиксирует сообщения, позволяет продавцу выдать оплаченные данные, а покупателю подтвердить получение, запросить возврат или открыть спор."
    )
    writer.paragraph(
        "Покупка выполняется внутри prisma.$transaction. Сервер проверяет наличие товара, доступность статуса LISTED, запрет покупки собственного товара и достаточность баланса при способе BALANCE. Затем товар атомарно переводится в SOLD, создаются Order и Payment, а для оплаты с баланса создаётся WalletTransaction типа PURCHASE_HOLD."
    )
    writer.caption("Листинг 3. Фрагмент транзакции покупки товара")
    writer.code(
        """
        const reserved = await tx.item.updateMany({
          where: { id: item.id, status: "LISTED" },
          data: { status: "SOLD", ownerId: buyer.id },
        });

        const order = await tx.order.create({
          data: {
            itemId: item.id,
            buyerId: buyer.id,
            sellerId: item.sellerId,
            price,
            status: "PAID",
          },
        });

        const payment = await tx.payment.create({
          data: {
            orderId: order.id,
            buyerId: buyer.id,
            sellerId: item.sellerId,
            method: paymentMethod,
            status: "PAID",
            amount: price,
            provider: paymentProvider(paymentMethod),
          },
        });
        """
    )
    writer.paragraph(
        "После оплаты продавец может сохранить deliveryData через PATCH /api/orders/:orderId/delivery. Это добавляет сообщение типа DELIVERY и системное сообщение о выдаче данных. "
        "Покупатель подтверждает получение через POST /api/orders/:orderId/confirm; в этом случае сделка становится COMPLETED, платёж получает статус RELEASED, а продавцу начисляется сумма через WalletTransaction типа SALE_RELEASE."
    )
    writer.paragraph(
        "Возврат реализован через POST /api/orders/:orderId/refund. Если сделка ещё не завершена, сервер переводит её в CANCELLED, возвращает товар в LISTED, очищает ownerId, меняет статус платежа на REFUNDED и при оплате с баланса возвращает средства покупателю. "
        "Если продавец уже выдал данные, возврат возможен только после открытия спора, что защищает продавца от односторонней отмены после получения цифрового товара."
    )

    writer.heading("3.4.4. Модуль отзывов, рейтинга и профиля продавца", 3)
    writer.paragraph(
        "Отзыв создаётся или обновляется через POST /api/orders/:orderId/review. Сервер разрешает это действие только покупателю и только для сделки со статусом COMPLETED. "
        "Ограничение unique по orderId не позволяет создать несколько отзывов по одной сделке, а upsert позволяет покупателю исправить уже оставленную оценку."
    )
    writer.paragraph(
        "Публичный профиль продавца реализован в users.ts. Эндпоинт GET /api/users/:userId возвращает данные пользователя, средний рейтинг, количество отзывов, активные товары и последние отзывы. "
        "Рейтинг также подмешивается в ответы каталога и карточки товара, поэтому покупатель видит надёжность продавца до покупки."
    )
    writer.paragraph(
        "Для связи до покупки реализованы DirectMessage и маршруты GET/POST /api/users/:userId/messages. Сервер запрещает отправлять сообщение самому себе и проверяет существование адресата."
    )

    writer.heading("3.4.5. Клиентские модули интерфейса", 3)
    writer.paragraph(
        "На клиенте каждая ключевая функция API имеет отдельную страницу или компонент. Каталог и карточки товара реализованы через Home, ItemCard и ItemDetail; публикация и редактирование - через AddItem и EditItem; операции пользователя - через Profile; сделки - через Deals и DealDetail; профиль продавца - через SellerProfile."
    )
    writer.caption("Листинг 4. Клиентский вызов создания сделки с выбранным способом оплаты")
    writer.code(
        """
        export async function buyItem(
          token: string,
          itemId: string,
          paymentMethod: PaymentMethod = "BALANCE"
        ) {
          return request<Deal>(`${API_URL}/orders/purchase/${itemId}`, token, {
            method: "POST",
            body: JSON.stringify({ paymentMethod }),
          });
        }
        """
    )
    writer.paragraph(
        "На странице ItemDetail пользователь выбирает способ оплаты: внутренний баланс сайта, банковскую карту или СБП. При оплате с баланса интерфейс проверяет достаточность средств и при необходимости автоматически переключает вариант оплаты. "
        "После успешной покупки вызывается refreshMe, чтобы обновить баланс в AuthContext, и пользователь переводится на страницу созданной сделки."
    )
    writer.paragraph(
        "Страница DealDetail объединяет несколько сценариев: просмотр товара и платежа, выдачу цифровых данных продавцом, чат сделки, открытие спора, подтверждение получения, возврат и отзыв. "
        "Состояния PAID, COMPLETED, DISPUTED и CANCELLED отображаются пользователю отдельными статусами."
    )

    writer.heading("3.4.6. Общие механизмы устойчивости", 3)
    writer.paragraph(
        "Устойчивость реализации обеспечивается несколькими общими механизмами: zod-схемы ограничивают входные данные, Prisma-транзакции защищают критичные операции покупки и возврата, middleware requireAuth закрывает приватные действия, а errorHandler возвращает единый JSON-ответ при ошибке. "
        "Файл env.ts централизует настройки PORT, CORS_ORIGIN, JWT_SECRET и DATABASE_URL."
    )

    writer.heading("3.5. Тестирование системы", 2)
    writer.paragraph(
        "Тестирование выполнялось по пользовательским сценариям и по негативным проверкам доступа. Основное внимание уделялось тем операциям, где меняются статусы сделок, остаток баланса и права участников."
    )
    writer.table(
        [
            ["Группа проверок", "Проверяемые сценарии", "Ожидаемый результат"],
            ["Auth", "Регистрация, повтор email/username, вход, GET /me, PATCH /me", "Пользователь получает токен, дубликаты отклоняются, данные профиля обновляются только у владельца."],
            ["Items", "Каталог, фильтрация, карточка, создание с изображением, PATCH с заменой/удалением картинки, удаление", "Товары корректно отображаются, изображения доступны, чужие и проданные товары нельзя менять."],
            ["Orders", "Покупка BALANCE/CARD/SBP, история покупок/продаж, чат, deliveryData, confirm, refund, dispute", "Статусы PAID/COMPLETED/CANCELLED/DISPUTED меняются по правилам, баланс и платежи фиксируются корректно."],
            ["Reviews/Users", "Отзыв после завершения, рейтинг продавца, публичный профиль, личное сообщение", "Оценка влияет на рейтинг, профиль показывает активные товары и отзывы, чат доступен только авторизованному пользователю."],
            ["Безопасность", "Запросы без токена, доступ к чужой сделке, покупка своего товара, недостаточный баланс, сообщение самому себе", "Сервер возвращает отказ и не изменяет данные."],
        ],
        [1.35, 3.1, 2.05],
    )
    writer.paragraph(
        "Проверки подтвердили, что реализованный прототип выполняет основной цикл работы маркетплейса: публикация товара, выбор в каталоге, оплата, создание сделки, передача оплаченных данных, переписка, подтверждение, начисление средств продавцу и формирование репутации через отзыв."
    )

    writer.heading("3.6. Результаты и перспективы развития", 2)
    writer.heading("3.6.1. Полученные результаты", 3)
    writer.paragraph(
        "В результате реализации получен рабочий прототип маркетплейса цифровых товаров. Реализованы регистрация и вход, каталог с категориями и поиском, публикация товаров с изображениями, карточка товара, внутренняя модель сделок и платежей, история покупок и продаж, чат сделки, выдача оплаченных данных, подтверждение получения, возврат, спор, отзывы, рейтинг продавца, публичный профиль продавца, личная переписка, локализация ru/en и переключение валюты в интерфейсе."
    )

    writer.heading("3.6.2. Ограничения текущей версии", 3)
    writer.paragraph(
        "Текущая версия является прототипом. Оплата картой и через СБП реализована как демонстрационная логика без подключения реального платёжного шлюза; административная панель для арбитража и модерации отсутствует; спор фиксируется статусом и перепиской, но решение администратора пока не реализовано; чаты работают без push-уведомлений и режима реального времени; отсутствуют расширенные инфраструктурные механизмы мониторинга, резервного копирования и нагрузочного тестирования."
    )

    writer.heading("3.6.3. Перспективы развития", 3)
    writer.paragraph(
        "Дальнейшее развитие проекта предполагает подключение реальных платёжных систем, расширение арбитража и административной панели, модерацию товаров и отзывов, развитие поиска и сортировки по рейтингу, добавление уведомлений и real-time обмена сообщениями, внедрение журналирования действий, CI/CD, мониторинга и резервного копирования."
    )
    writer.paragraph(
        "Таким образом, реализованная система демонстрирует завершённый базовый цикл взаимодействия пользователей маркетплейса цифровых товаров и создаёт основу для дальнейшего промышленного развития."
    )


def build_chapter_preserving_structure(writer: ChapterWriter) -> None:
    def paras(*texts: str) -> None:
        for text in texts:
            writer.paragraph(text)

    writer.heading("Глава 3. Реализация", 1)

    writer.heading("3.1. Выбор технологий и инструментов", 2)
    paras(
        "При реализации маркетплейса цифровых товаров был использован современный стек веб-технологий, обеспечивающий удобство разработки, масштабируемость и соответствие функциональным и нефункциональным требованиям, сформулированным в предыдущей главе.",
        "В качестве технологии клиентской части был выбран React с использованием TypeScript. Данный выбор обусловлен необходимостью создания динамического пользовательского интерфейса, поддержки компонентного подхода и строгой типизации, снижающей количество ошибок на этапе разработки. React позволяет эффективно реализовывать сложные пользовательские сценарии, такие как работа с каталогом товаров, карточкой товара, сделками и пользовательским профилем, без перезагрузки страниц. Для сборки клиентской части используется Vite, а маршрутизация между экранами реализована через React Router.",
        "Серверная часть реализована на платформе Node.js с использованием Express и TypeScript, что обеспечивает единый язык программирования для клиентской и серверной логики. Использование серверного JavaScript упрощает поддержку проекта и ускоряет разработку REST API, необходимого для взаимодействия клиентской части с базой данных и бизнес-логикой приложения.",
        "В качестве системы управления базами данных была выбрана PostgreSQL. Данная СУБД обеспечивает надёжное хранение структурированных данных, поддержку связей между сущностями и транзакционность, что является критически важным для обработки заказов, пользователей, цифровых товаров, платежей и операций внутреннего баланса.",
        "Для работы с базой данных используется Prisma ORM. В проекте также применяются JWT для аутентификации, bcryptjs для хэширования паролей, zod для валидации входных данных и multer для загрузки изображений товаров.",
        "Выбранный набор технологий обеспечивает баланс между производительностью, надёжностью и удобством сопровождения системы, а также соответствует архитектуре и функциональности реализованного проекта маркетплейса цифровых товаров.",
    )

    writer.heading("3.2. Архитектура программной системы", 2)
    paras(
        "Реализованная система имеет классическую клиент-серверную архитектуру и разделена на два независимых приложения в одном репозитории:",
        "client/ - клиентская часть (React + TypeScript, Vite), отвечает за интерфейс, навигацию и вызовы API.",
        "server/ - серверная часть (Node.js + Express + TypeScript), отвечает за бизнес-логику, авторизацию и работу с БД.",
        "Такое разделение позволяет независимо развивать пользовательский интерфейс и серверное API, при этом вся бизнес-логика, связанная с доступом к данным, статусами сделок и финансовыми операциями, остаётся на стороне backend.",
    )

    writer.heading("3.2.1. Общая схема взаимодействия компонентов", 3)
    paras(
        "Поток данных организован следующим образом:",
        "Frontend => REST API => База данных:",
        "Frontend отправляет запросы на сервер по HTTP к API с базовым префиксом /api (пример: /api/items, /api/auth).",
        "Backend обрабатывает запросы, применяет проверки и бизнес-правила, после чего обращается к базе данных через ORM.",
        "База данных хранит пользователей, товары, сделки, платежи, операции кошелька, сообщения, отзывы и личную переписку; сервер возвращает клиенту данные в формате JSON, а для изображений товаров - бинарные данные.",
        "Например, при покупке товара страница карточки вызывает клиентскую функцию buyItem, сервер создаёт сделку, платёж и системное сообщение в одной транзакции, после чего интерфейс перенаправляет пользователя на страницу сделки.",
    )

    writer.heading("3.2.2. Серверное приложение и основные модули", 3)
    paras(
        "Backend построен как набор модулей (роутов) и middleware:",
        "Маршруты API:",
        "GET /api/health - проверка доступности серверного приложения.",
        "POST /api/auth/register - регистрация.",
        "POST /api/auth/login - вход.",
        "GET /api/auth/me - получение данных текущего пользователя.",
        "PATCH /api/auth/me - изменение username текущего пользователя.",
        "GET /api/items - список товаров (каталог).",
        "GET /api/items/:id - карточка товара.",
        "GET /api/items/mine - товары текущего продавца.",
        "GET /api/items/:id/image - получение изображения товара.",
        "POST /api/items - создание товара (с поддержкой загрузки изображения).",
        "PUT /api/items/:id - редактирование товара через JSON.",
        "PATCH /api/items/:id - частичное редактирование товара, включая замену или удаление изображения.",
        "DELETE /api/items/:id - удаление товара.",
        "POST /api/orders/purchase/:itemId - покупка товара и создание сделки со статусом PAID.",
        "GET /api/orders/my - история покупок.",
        "GET /api/orders/sales - история продаж.",
        "GET /api/orders/:orderId - детальная информация о сделке.",
        "POST /api/orders/:orderId/messages - отправка сообщения в чат сделки.",
        "PATCH /api/orders/:orderId/delivery - сохранение оплаченных данных продавцом.",
        "POST /api/orders/:orderId/confirm - подтверждение получения товара покупателем.",
        "POST /api/orders/:orderId/refund - отмена сделки и возврат средств по допустимому сценарию.",
        "POST /api/orders/:orderId/dispute - открытие спора по сделке.",
        "POST /api/orders/:orderId/review - создание или обновление отзыва по завершённой сделке.",
        "GET /api/users/:userId - публичный профиль продавца.",
        "GET /api/users/:userId/messages - получение личной переписки с продавцом.",
        "POST /api/users/:userId/messages - отправка личного сообщения продавцу.",
        "Авторизация и доступ:",
        "защищённые эндпоинты используют middleware requireAuth;",
        "аутентификация реализована через JWT в заголовке Authorization: Bearer <token>.",
        "Обработка ошибок:",
        "централизованная обработка ошибок вынесена в middleware errorHandler, что упрощает сопровождение и единообразие ответов.",
    )

    writer.heading("3.2.3. Хранение данных и модель предметной области", 3)
    paras(
        "Для доступа к данным используется PostgreSQL через Prisma. В базе данных выделены ключевые сущности:",
        "User - пользователь (email, username, passwordHash, role, balance), связанный с товарами, покупками, продажами, отзывами, платежами, операциями кошелька и сообщениями.",
        "Item - товар (title, description, price, status, category, imageData, imageMime, sellerId, ownerId).",
        "Order - сделка/заказ (itemId, buyerId, sellerId, price, status, deliveryData, confirmedAt, disputedAt, cancelledAt, createdAt).",
        "Payment - платёж по сделке (orderId, buyerId, sellerId, method, status, amount, provider, providerPaymentId, releasedAt, refundedAt).",
        "WalletTransaction - операция внутреннего баланса пользователя (userId, orderId, paymentId, type, amount, balanceAfter, note).",
        "OrderMessage - сообщения внутри сделки (orderId, authorId, body, type, createdAt).",
        "Review - отзыв покупателя о продавце по завершённой сделке (orderId, buyerId, sellerId, rating, comment).",
        "DirectMessage - личные сообщения между пользователем и продавцом (senderId, recipientId, body, createdAt).",
        "Изображения товаров реализованы без отдельного файлового хранилища: картинка хранится в БД как imageData (Bytes) и imageMime, а отдача выполняется отдельным эндпоинтом GET /api/items/:id/image.",
    )

    writer.heading("3.2.4. Ключевое бизнес-правило покупки", 3)
    paras(
        "Операция покупки реализована как транзакция:",
        "проверка доступности товара (статус LISTED);",
        "запрет покупки собственного товара;",
        "выбор способа оплаты: внутренний баланс сайта, карта или СБП;",
        "проверка баланса покупателя при оплате способом BALANCE;",
        "списание средств у покупателя и создание операции PURCHASE_HOLD при оплате с баланса;",
        "перевод товара в статус SOLD и назначение владельца;",
        "создание записи заказа и записи Payment со статусом PAID;",
        "создание системного сообщения в чате сделки;",
        "после подтверждения получения сделка переводится в COMPLETED, платёж - в RELEASED, а продавцу начисляются средства через операцию SALE_RELEASE;",
        "при допустимом возврате сделка переводится в CANCELLED, товар возвращается в LISTED, платёж получает статус REFUNDED, а покупателю возвращаются средства;",
        "при конфликте сделка может быть переведена в статус DISPUTED.",
        "Такой подход обеспечивает согласованность данных при оформлении сделки и снижает риск спорных ситуаций, так как продавец получает средства только после подтверждения получения цифрового товара покупателем.",
    )

    writer.heading("3.3. Проектирование базы данных", 2)
    paras(
        "База данных спроектирована для хранения пользователей, цифровых товаров, сделок, платежей, операций кошелька, сообщений, отзывов и личной переписки. В проекте используется реляционная модель данных (PostgreSQL) и ORM Prisma; структура описана в виде сущностей User, Item, Order, Payment, WalletTransaction, OrderMessage, Review, DirectMessage, а также перечислений Role, ItemStatus, ItemCategory, OrderStatus, PaymentMethod, PaymentStatus, WalletTransactionType и OrderMessageType."
    )

    writer.heading("3.3.1. ER-модель и связи между сущностями", 3)
    writer.caption("Рисунок 1. ER-модель базы данных маркетплейса цифровых товаров.")
    writer.picture(ER_IMAGE)
    paras(
        "Основные связи:",
        "User (продавец) 1 - N Item: один пользователь может разместить много товаров; у товара один продавец (sellerId).",
        "User (владелец) 1 - N Item: после покупки товар получает владельца; у пользователя может быть много купленных товаров (ownerId у товара может быть пустым до покупки).",
        "User (покупатель) 1 - N Order: пользователь может иметь много покупок (buyerId).",
        "User (продавец) 1 - N Order: пользователь может иметь много продаж (sellerId).",
        "Item 1 - N Order на уровне схемы: история заказов связана с товаром через itemId; при этом бизнес-логика переводит товар в SOLD и не позволяет повторно купить уже проданный товар.",
        "Order 1 - N OrderMessage: в рамках одной сделки может быть несколько сообщений чата.",
        "Order 1 - 1 Review: по одной завершённой сделке покупатель может оставить один отзыв.",
        "Order 1 - 1 Payment: платёж фиксирует метод, сумму, статус и сведения о провайдере.",
        "Payment 1 - N WalletTransaction и User 1 - N WalletTransaction: операции кошелька фиксируют резервирование, возврат и зачисление средств.",
        "User (покупатель) 1 - N Review и User (продавец) 1 - N Review: отзывы связывают покупателя и продавца и используются для расчёта рейтинга.",
        "User 1 - N DirectMessage: пользователь может отправлять и получать личные сообщения, связанные с профилем продавца.",
    )

    writer.heading("3.3.2. Описание сущностей и полей", 3)
    writer.paragraph("Ниже приведены таблицы (модели) и их поля в том виде, как они реализованы в проекте.")
    writer.caption("Таблица 1. Таблица User (Пользователи)")
    writer.table(
        [
            ["Поле", "Тип", "Ограничения/назначение"],
            ["id", "String", "PK, генерируется автоматически"],
            ["email", "String", "UNIQUE, email пользователя"],
            ["username", "String", "UNIQUE, имя пользователя"],
            ["passwordHash", "String", "хэш пароля"],
            ["role", "Enum (Role)", "роль (USER/ADMIN), по умолчанию USER"],
            ["balance", "Int", "баланс, по умолчанию 1000"],
            ["createdAt", "DateTime", "дата создания"],
            ["updatedAt", "DateTime", "дата обновления"],
        ],
        [1.6, 1.7, 3.0],
    )
    paras(
        "Связанные данные:",
        "items - товары, которые пользователь разместил как продавец (1-N).",
        "ownedItems - товары, которые пользователь купил (1-N).",
        "orders - покупки пользователя (1-N).",
        "sales - продажи пользователя (1-N).",
        "reviewsGiven - отзывы, оставленные пользователем как покупателем (1-N).",
        "reviewsReceived - отзывы, полученные пользователем как продавцом (1-N).",
        "sentDirectMessages и receivedDirectMessages - личные сообщения пользователя (1-N).",
        "buyerPayments и sellerPayments - платежи, где пользователь выступает покупателем или продавцом.",
        "walletTransactions - операции внутреннего баланса пользователя.",
    )

    writer.caption("Таблица 2. Таблица Item (Товары)")
    writer.table(
        [
            ["Поле", "Тип", "Ограничения/назначение"],
            ["id", "String", "PK, генерируется автоматически"],
            ["title", "String", "название товара"],
            ["description", "String", "описание товара"],
            ["price", "Decimal (10,2)", "цена товара"],
            ["status", "Enum (ItemStatus)", "LISTED/SOLD, по умолчанию LISTED"],
            ["category", "Enum (ItemCategory)", "категория, по умолчанию OTHER"],
            ["imageData", "Bytes (nullable)", "бинарные данные изображения, если загружено"],
            ["imageMime", "String (nullable)", "MIME-тип изображения"],
            ["sellerId", "String", "FK на User.id, продавец товара"],
            ["ownerId", "String (nullable)", "FK на User.id, владелец после покупки"],
            ["createdAt", "DateTime", "дата создания"],
            ["updatedAt", "DateTime", "дата обновления"],
        ],
        [1.6, 1.7, 3.0],
    )
    paras("Смысл полей sellerId и ownerId:", "sellerId заполняется всегда и определяет автора размещения товара.", "ownerId заполняется только после покупки (до этого NULL).")

    writer.caption("Таблица 3. Таблица Order (Заказы)")
    writer.table(
        [
            ["Поле", "Тип", "Ограничения/назначение"],
            ["id", "String", "PK, генерируется автоматически"],
            ["itemId", "String", "FK на Item.id, товар сделки"],
            ["buyerId", "String", "FK на User.id, покупатель"],
            ["sellerId", "String", "FK на User.id, продавец"],
            ["price", "Decimal (10,2)", "итоговая цена сделки"],
            ["status", "Enum (OrderStatus)", "PAID, COMPLETED, DISPUTED, CANCELLED"],
            ["deliveryData", "String (nullable)", "оплаченные данные, выданные продавцом"],
            ["confirmedAt", "DateTime (nullable)", "дата подтверждения получения"],
            ["disputedAt", "DateTime (nullable)", "дата открытия спора"],
            ["cancelledAt", "DateTime (nullable)", "дата отмены сделки"],
            ["createdAt", "DateTime", "дата создания заказа"],
        ],
        [1.6, 1.7, 3.0],
    )
    paras(
        "Сущность OrderMessage (Сообщения сделки)",
        "Сообщение сделки хранит идентификатор сделки, автора, текст сообщения, тип сообщения (TEXT, SYSTEM, DELIVERY) и дату создания. Данная сущность используется для переписки между покупателем и продавцом и фиксации системных событий.",
        "Сущность Review (Отзывы)",
        "Отзыв содержит ссылку на сделку, покупателя, продавца, пятизвёздочную оценку, комментарий и даты создания/обновления. Ограничение уникальности orderId не позволяет создать несколько отзывов по одной сделке.",
        "Сущность DirectMessage (Личные сообщения)",
        "Личное сообщение хранит отправителя, получателя, текст сообщения и дату создания. Используется для общения покупателя с продавцом через публичный профиль продавца.",
    )

    writer.caption("Таблица 4. Таблица Payment (Платежи)")
    writer.table(
        [
            ["Поле", "Тип", "Ограничения/назначение"],
            ["id", "String", "PK, генерируется автоматически"],
            ["orderId", "String", "FK на Order.id, UNIQUE"],
            ["buyerId", "String", "FK на User.id, покупатель"],
            ["sellerId", "String", "FK на User.id, продавец"],
            ["method", "Enum (PaymentMethod)", "BALANCE, CARD, SBP"],
            ["status", "Enum (PaymentStatus)", "PENDING, PAID, RELEASED, REFUNDED, FAILED"],
            ["amount", "Decimal (10,2)", "сумма платежа"],
            ["provider", "String (nullable)", "условный провайдер платежа"],
            ["providerPaymentId", "String (nullable)", "уникальный идентификатор платежа у провайдера"],
            ["releasedAt", "DateTime (nullable)", "дата перевода средств продавцу"],
            ["refundedAt", "DateTime (nullable)", "дата возврата средств"],
        ],
        [1.6, 1.7, 3.0],
    )

    writer.caption("Таблица 5. Таблица WalletTransaction (Операции баланса)")
    writer.table(
        [
            ["Поле", "Тип", "Ограничения/назначение"],
            ["id", "String", "PK, генерируется автоматически"],
            ["userId", "String", "FK на User.id"],
            ["orderId", "String (nullable)", "связь с заказом"],
            ["paymentId", "String (nullable)", "связь с платежом"],
            ["type", "Enum (WalletTransactionType)", "PURCHASE_HOLD, REFUND, SALE_RELEASE"],
            ["amount", "Decimal (10,2)", "сумма изменения баланса"],
            ["balanceAfter", "Int", "баланс пользователя после операции"],
            ["note", "String (nullable)", "пояснение к операции"],
            ["createdAt", "DateTime", "дата создания"],
        ],
        [1.6, 1.7, 3.0],
    )
    writer.caption("Листинг 1. Фрагмент схемы Prisma для платежей и кошелька")
    writer.code(
        """
        enum PaymentMethod {
          BALANCE
          CARD
          SBP
        }

        enum PaymentStatus {
          PENDING
          PAID
          RELEASED
          REFUNDED
          FAILED
        }

        enum WalletTransactionType {
          PURCHASE_HOLD
          REFUND
          SALE_RELEASE
        }

        model Payment {
          id                String        @id @default(cuid())
          orderId           String        @unique
          buyerId           String
          sellerId          String
          method            PaymentMethod
          status            PaymentStatus @default(PENDING)
          amount            Decimal       @db.Decimal(10, 2)
          provider          String?
          providerPaymentId String?       @unique
          releasedAt        DateTime?
          refundedAt        DateTime?
        }
        """
    )

    writer.heading("3.3.3. Перечисления", 3)
    paras(
        "Role: USER, ADMIN.",
        "ItemStatus: LISTED, SOLD.",
        "ItemCategory: ACCOUNTS, KEYS, SUBSCRIPTIONS, SERVICES, GAME_CURRENCIES, NFT_TOKENS, OTHER.",
        "OrderStatus: PAID, COMPLETED, DISPUTED, CANCELLED.",
        "PaymentMethod: BALANCE, CARD, SBP.",
        "PaymentStatus: PENDING, PAID, RELEASED, REFUNDED, FAILED.",
        "WalletTransactionType: PURCHASE_HOLD, REFUND, SALE_RELEASE.",
        "OrderMessageType: TEXT, SYSTEM, DELIVERY.",
    )

    writer.heading("3.4. Реализация функциональных модулей", 2)
    writer.paragraph("В реализованном проекте функциональность системы разделена на серверные модули (REST API) и клиентские экраны/компоненты.")

    writer.heading("3.4.1. Модуль регистрации и аутентификации (Auth)", 3)
    paras(
        "Назначение: управление доступом пользователей к системе.",
        "Реализовано (API):",
        "POST /api/auth/register - регистрация (email, username, password) с валидацией входных данных.",
        "POST /api/auth/login - вход по email/паролю.",
        "GET /api/auth/me - получение данных текущего пользователя (по токену).",
        "PATCH /api/auth/me - изменение username с проверкой уникальности.",
        "Ключевые особенности:",
        "пароли хранятся в виде хэша (bcrypt);",
        "используется JWT-токен сроком действия 7 дней;",
        "токен передаётся через Authorization: Bearer <token>;",
        "проверка авторизации вынесена в middleware requireAuth.",
    )
    writer.caption("Листинг 2. Проверка JWT-токена в middleware requireAuth")
    writer.code(
        """
        export const requireAuth = (req, res, next) => {
          const header = req.headers.authorization;
          if (!header?.startsWith("Bearer ")) {
            return res.status(401).json({ message: "Missing token" });
          }

          try {
            const payload = jwt.verify(header.slice(7), env.JWT_SECRET);
            req.user = payload;
            next();
          } catch {
            res.status(401).json({ message: "Invalid token" });
          }
        };
        """
    )

    writer.heading("3.4.2. Модуль каталога и карточек товаров (Items)", 3)
    paras(
        "Назначение: публикация, просмотр и поиск цифровых товаров.",
        "Реализовано (API):",
        "GET /api/items - получение списка товаров с параметрами:",
        "search - поиск по названию (без учёта регистра),",
        "category - фильтр по категории,",
        "status - фильтр по статусу (по умолчанию LISTED, доступно ALL).",
        "GET /api/items/:id - детальная информация о товаре.",
        "GET /api/items/:id/image - получение изображения товара (если загружено).",
        "Реализовано для продавца (API):",
        "GET /api/items/mine - список товаров текущего продавца.",
        "POST /api/items - добавление товара (multipart/form-data) с поддержкой изображения.",
        "PUT /api/items/:id - редактирование товара.",
        "PATCH /api/items/:id - multipart-редактирование с заменой или удалением изображения.",
        "DELETE /api/items/:id - удаление товара.",
        "Ключевые особенности:",
        "категории реализованы через enum (например: ACCOUNTS, KEYS, SUBSCRIPTIONS, SERVICES, GAME_CURRENCIES, NFT_TOKENS, OTHER);",
        "загрузка изображения реализована через multer, ограничения:",
        "размер до 5 MB,",
        "форматы: jpg/png/webp;",
        "изображение хранится в БД (imageData + imageMime), а наличие изображения отдаётся флагом hasImage;",
        "в ответах каталога и карточки товара дополнительно передаются сведения о продавце и его рейтинге, что позволяет отображать надёжность продавца непосредственно в интерфейсе;",
        "перед изменением или удалением сервер проверяет, что товар принадлежит текущему пользователю и ещё не продан.",
    )

    writer.heading("3.4.3. Модуль сделок, арбитража и истории операций (Orders)", 3)
    paras(
        "Назначение: оформление покупки, создание сделки, резервирование средств, передача оплаченных данных, переписка участников, подтверждение получения, возврат, открытие спора и формирование истории покупок/продаж.",
        "Реализовано (API):",
        "POST /api/orders/purchase/:itemId - покупка товара.",
        "GET /api/orders/my - история покупок пользователя.",
        "GET /api/orders/sales - история продаж пользователя.",
        "GET /api/orders/:orderId - получение детальной информации о сделке.",
        "POST /api/orders/:orderId/messages - отправка сообщения в чат сделки.",
        "PATCH /api/orders/:orderId/delivery - добавление продавцом оплаченных данных по товару.",
        "POST /api/orders/:orderId/confirm - подтверждение получения товара покупателем и перевод средств продавцу.",
        "POST /api/orders/:orderId/refund - возврат средств и отмена сделки.",
        "POST /api/orders/:orderId/dispute - открытие спора по сделке.",
        "Ключевые особенности покупки:",
        "запрет покупки собственного товара;",
        "проверка доступности товара (status = LISTED);",
        "поддержка способов оплаты BALANCE, CARD и SBP;",
        "проверка баланса покупателя при оплате способом BALANCE;",
        "операция покупки выполняется транзакционно:",
        "списание средств у покупателя при оплате с баланса,",
        "средства резервируются до подтверждения получения товара;",
        "перевод товара в SOLD и назначение владельца,",
        "создание записи заказа, платежа и системного сообщения.",
    )
    writer.caption("Листинг 3. Фрагмент транзакции покупки товара")
    writer.code(
        """
        const reserved = await tx.item.updateMany({
          where: { id: item.id, status: "LISTED" },
          data: { status: "SOLD", ownerId: buyer.id },
        });

        const createdOrder = await tx.order.create({
          data: {
            itemId: item.id,
            buyerId: buyer.id,
            sellerId: item.sellerId,
            price,
            status: "PAID",
          },
        });

        const payment = await tx.payment.create({
          data: {
            orderId: createdOrder.id,
            buyerId: buyer.id,
            sellerId: item.sellerId,
            method: paymentMethod,
            status: "PAID",
            amount: price,
            provider: paymentProvider(paymentMethod),
          },
        });
        """
    )
    paras(
        "После оплаты продавец может сохранить deliveryData через PATCH /api/orders/:orderId/delivery. Это добавляет сообщение типа DELIVERY и системное сообщение о выдаче данных.",
        "Покупатель подтверждает получение через POST /api/orders/:orderId/confirm; в этом случае сделка становится COMPLETED, платёж получает статус RELEASED, а продавцу начисляется сумма через WalletTransaction типа SALE_RELEASE.",
        "Возврат реализован через POST /api/orders/:orderId/refund. Если сделка ещё не завершена, сервер переводит её в CANCELLED, возвращает товар в LISTED, очищает ownerId, меняет статус платежа на REFUNDED и при оплате с баланса возвращает средства покупателю.",
        "Если продавец уже выдал данные, возврат возможен только после открытия спора, что защищает продавца от односторонней отмены после получения цифрового товара.",
    )

    writer.heading("3.4.4. Модуль отзывов, рейтинга и профиля продавца", 3)
    paras(
        "Назначение: формирование доверия между пользователями за счёт отзывов, рейтинга продавца и публичной страницы продавца.",
        "Реализовано (API):",
        "POST /api/orders/:orderId/review - создание или обновление отзыва по завершённой сделке.",
        "GET /api/users/:userId - публичный профиль продавца с рейтингом, активными товарами и отзывами.",
        "GET /api/users/:userId/messages - получение личной переписки с продавцом.",
        "POST /api/users/:userId/messages - отправка личного сообщения продавцу.",
        "Ключевые особенности:",
        "оценка продавца выставляется по пятизвёздочной шкале;",
        "отзыв может быть оставлен только покупателем после завершения сделки;",
        "рейтинг продавца рассчитывается на основе полученных отзывов и отображается в карточке товара, на странице товара и в публичном профиле продавца;",
        "личный чат позволяет задать продавцу вопрос до покупки или вне конкретной сделки.",
    )

    writer.heading("3.4.5. Клиентские модули интерфейса (Frontend)", 3)
    paras(
        "Назначение: пользовательские сценарии через веб-интерфейс.",
        "Реализованы основные экраны:",
        "Главная/каталог - просмотр списка товаров, поиск и фильтрация;",
        "Карточка товара - просмотр деталей товара, выбор способа оплаты и покупка;",
        "Регистрация / Вход - управление сессией пользователя;",
        "Профиль - баланс, статистика покупок/продаж/товаров, быстрые действия и настройки;",
        "Раздел Сделки - список покупок и продаж с переходом к конкретной сделке;",
        "Страница сделки - сведения о товаре и платеже, оплаченные данные, чат, подтверждение получения, возврат, открытие спора и отзыв;",
        "Публичный профиль продавца - рейтинг, отзывы, активные товары и личный чат;",
        "Добавление и редактирование товара - формы управления карточкой товара и изображением.",
        "Дополнительно реализованы:",
        "Контекст авторизации (AuthContext) - хранение токена и данных пользователя, автоматическая загрузка /api/auth/me, обновление баланса через refreshMe;",
        "Контекст настроек (SettingsContext) - переключение языка (ru/en) и валюты (RUB/USD) на уровне интерфейса.",
    )
    writer.caption("Листинг 4. Клиентский вызов создания сделки с выбранным способом оплаты")
    writer.code(
        """
        export async function buyItem(
          token: string,
          itemId: string,
          paymentMethod: PaymentMethod = "BALANCE"
        ) {
          return request<Deal>(`${API_URL}/orders/purchase/${itemId}`, token, {
            method: "POST",
            body: JSON.stringify({ paymentMethod }),
          });
        }
        """
    )
    paras(
        "На странице ItemDetail пользователь выбирает способ оплаты: внутренний баланс сайта, банковскую карту или СБП. При оплате с баланса интерфейс проверяет достаточность средств и при необходимости автоматически переключает вариант оплаты.",
        "После успешной покупки вызывается refreshMe, чтобы обновить баланс в AuthContext, и пользователь переводится на страницу созданной сделки.",
        "Страница DealDetail объединяет несколько сценариев: просмотр товара и платежа, выдачу цифровых данных продавцом, чат сделки, открытие спора, подтверждение получения, возврат и отзыв.",
        "Состояния PAID, COMPLETED, DISPUTED и CANCELLED отображаются пользователю отдельными статусами.",
    )

    writer.heading("3.4.6. Общие механизмы (для устойчивости модуляции)", 3)
    paras(
        "валидация входных данных на сервере выполнена через zod;",
        "единая обработка ошибок вынесена в middleware errorHandler;",
        "доступ к защищённым операциям ограничен middleware requireAuth;",
        "критичные операции покупки, подтверждения и возврата выполняются через prisma.$transaction;",
        "файл env.ts централизует настройки PORT, CORS_ORIGIN, JWT_SECRET и DATABASE_URL.",
    )

    writer.heading("3.5. Тестирование системы", 2)
    writer.paragraph(
        "Тестирование выполнено для проверки корректности работы реализованных модулей маркетплейса: аутентификация пользователей, управление товарами, совершение сделок, резервирование средств, чат сделки, подтверждение получения, возврат, отзывы, рейтинг продавца и публичный профиль продавца. Проверка проводилась по позитивным и негативным сценариям, включая контроль прав доступа и корректность изменения статусов."
    )

    writer.heading("3.5.1. Подход к тестированию", 3)
    paras(
        "В ходе тестирования проверялись:",
        "корректность бизнес-логики (покупка товара, создание сделки, резервирование средств, подтверждение получения, перевод средств продавцу, возврат, открытие спора);",
        "валидация данных (обязательные поля, формат email, ограничения на цену, категории);",
        "контроль доступа (запрет действий без авторизации, запрет работы с чужими данными);",
        "обработка ошибок (ожидаемые HTTP-коды и сообщения).",
    )

    writer.heading("3.5.2. Функциональные тесты (реализованные сценарии)", 3)
    paras(
        "1) Регистрация и вход (Auth)",
        "регистрация нового пользователя (POST /api/auth/register);",
        "запрет регистрации при повторяющемся email/username (ожидается конфликт);",
        "вход по email/паролю (POST /api/auth/login);",
        "получение данных текущего пользователя по токену (GET /api/auth/me);",
        "изменение username (PATCH /api/auth/me) с проверкой уникальности.",
        "2) Каталог и товары (Items)",
        "получение списка товаров (GET /api/items) с фильтрами search/category/status;",
        "получение карточки товара (GET /api/items/:id);",
        "просмотр списка своих товаров продавца (GET /api/items/mine);",
        "создание товара с изображением и без изображения (POST /api/items);",
        "редактирование товара:",
        "JSON-обновление (PUT /api/items/:id);",
        "multipart-обновление с заменой/удалением картинки (PATCH /api/items/:id);",
        "удаление товара (DELETE /api/items/:id);",
        "получение изображения товара (GET /api/items/:id/image).",
        "3) Сделки, чат и история (Orders)",
        "покупка товара (POST /api/orders/purchase/:itemId) с оплатой BALANCE, CARD и SBP;",
        "проверка корректного изменения статуса товара на SOLD и назначения владельца;",
        "проверка списания средств у покупателя и отсутствия начисления продавцу до подтверждения получения;",
        "проверка создания Payment и WalletTransaction при оплате с баланса;",
        "проверка добавления продавцом оплаченных данных по сделке;",
        "проверка отправки сообщений в чат сделки;",
        "проверка подтверждения получения товара покупателем и последующего начисления средств продавцу;",
        "проверка возврата средств и отмены сделки через POST /api/orders/:orderId/refund;",
        "проверка открытия спора и изменения статуса сделки на DISPUTED;",
        "проверка создания отзыва по завершённой сделке и обновления рейтинга продавца;",
        "проверка просмотра публичного профиля продавца и отправки личного сообщения.",
        "просмотр истории покупок (GET /api/orders/my);",
        "просмотр истории продаж (GET /api/orders/sales).",
    )

    writer.heading("3.5.3. Негативные сценарии и проверки безопасности доступа", 3)
    paras(
        "Проверены типовые ошибки и ограничения:",
        "попытка обращения к защищённым эндпоинтам без токена => отказ в доступе;",
        "попытка редактировать/удалять товар другого пользователя => запрет операции;",
        "запрет изменения и удаления проданного товара;",
        "запрет покупки собственного товара;",
        "запрет покупки при недостаточном балансе;",
        "запрет доступа к чужой сделке и чужому чату сделки;",
        "запрет подтверждения получения продавцом вместо покупателя;",
        "запрет добавления оплаченных данных покупателем вместо продавца;",
        "запрет возврата уже завершённой сделки;",
        "запрет возврата после выдачи данных без предварительного открытия спора;",
        "запрет создания отзыва до завершения сделки и запрет отзыва от лица продавца;",
        "запрет отправки личного сообщения самому себе;",
        "обработка некорректной категории товара (ошибка валидации);",
        "ограничения на загрузку изображений: допустимые MIME-типы и лимит размера файла.",
    )

    writer.heading("3.5.4. Результаты", 3)
    writer.paragraph(
        "По результатам тестирования подтверждена работоспособность ключевых пользовательских сценариев: регистрация/вход, публикация и просмотр товаров, покупка, создание сделки, передача оплаченных данных, переписка участников, подтверждение получения, возврат, отзывы и просмотр публичного профиля продавца. Также подтверждены ограничения прав доступа, корректность изменения статусов и валидация входных данных."
    )

    writer.heading("3.6. Результаты и перспективы развития", 2)
    writer.paragraph(
        "В ходе практики был реализован программный прототип веб-приложения «Маркетплейс цифровых товаров», включающий клиентскую и серверную части, а также базу данных."
    )

    writer.heading("3.6.1. Полученные результаты", 3)
    paras(
        "В рамках реализованного проекта выполнено следующее:",
        "реализована регистрация и аутентификация пользователей с использованием JWT и хранением паролей в виде хэша;",
        "реализован каталог цифровых товаров с поиском и фильтрацией по параметрам;",
        "реализовано управление товарами для продавца: создание, редактирование, удаление, загрузка изображения;",
        "реализована покупка товара с выбором способа оплаты (BALANCE, CARD, SBP), проверками бизнес-правил, созданием сделки, платежа и резервированием средств до подтверждения получения;",
        "реализованы операции внутреннего баланса: резервирование средств покупателя, возврат и зачисление продавцу;",
        "реализованы разделы истории операций: покупки и продажи;",
        "реализованы отдельный раздел сделок и страница сделки с оплаченными данными, чатом, подтверждением получения, возвратом и открытием спора;",
        "реализована пятизвёздочная система отзывов с комментариями по завершённым сделкам;",
        "реализован рейтинг продавца и его отображение в карточке товара, на странице товара и в публичном профиле продавца;",
        "реализован публичный профиль продавца со списком активных товаров, отзывами и личным чатом;",
        "реализована базовая локализация интерфейса (ru/en) и переключение валюты на уровне UI (RUB/USD);",
        "проведено функциональное тестирование основных модулей и негативных сценариев.",
    )

    writer.heading("3.6.2. Ограничения текущей версии", 3)
    paras(
        "Текущая версия проекта имеет ограничения, связанные с уровнем готовности прототипа:",
        "оплата картой и через СБП реализована как демонстрационная логика без подключения реального платёжного шлюза;",
        "внутренний баланс используется для демонстрации резервирования, возврата и зачисления средств;",
        "отсутствует полноценная административная панель для ручного рассмотрения спорных ситуаций;",
        "механизм спора реализован на уровне статуса сделки и переписки, однако сценарий вынесения решения администратором требует дальнейшей доработки;",
        "отсутствуют инструменты модерации отзывов, товаров и сообщений;",
        "чат реализован в рамках сделки и профиля продавца, однако отсутствуют push-уведомления и режим обмена сообщениями в реальном времени;",
        "отсутствуют расширенные нефункциональные проверки (нагрузочное тестирование, мониторинг, резервное копирование на уровне инфраструктуры).",
    )

    writer.heading("3.6.3. Перспективы развития", 3)
    paras(
        "В качестве направлений развития системы выделены:",
        "интеграция реальных платёжных систем и построение полного процесса оплаты и возвратов;",
        "развитие системы отзывов и рейтингов: модерация комментариев, фильтрация по оценке, отображение динамики рейтинга продавца;",
        "разработка полноценной административной панели арбитража с просмотром спорных сделок, переписки, оплаченных данных и вынесением решения по переводу или возврату средств;",
        "развитие модели безопасности (защита от перебора, лимиты запросов, аудит действий);",
        "реализация административной панели для модерации товаров и пользователей;",
        "расширение поисковой системы (сортировки, фильтры по рейтингу/популярности);",
        "развитие чатов: уведомления о новых сообщениях, отметки прочтения, вложения и обмен сообщениями в реальном времени;",
        "подготовка к промышленной эксплуатации: логирование, мониторинг, CI/CD, резервное копирование.",
        "В результате выполненной работы получен реализованный прототип маркетплейса цифровых товаров, демонстрирующий основной цикл взаимодействия пользователей (публикация - выбор - покупка - фиксация сделки - передача данных - подтверждение - отзыв) и имеющий понятные направления для дальнейшего расширения функциональности.",
    )


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    make_er_diagram(ER_IMAGE)

    doc = Document(SOURCE)
    anchor = remove_old_chapter(doc)
    writer = ChapterWriter(doc, anchor)
    build_chapter_preserving_structure(writer)

    doc.save(OUT_DOCX)
    print(OUT_DOCX)


if __name__ == "__main__":
    main()
